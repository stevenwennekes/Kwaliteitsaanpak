"""Docx builder."""

import pathlib
import shutil
from docx import Document
from docx.enum.text import WD_COLOR_INDEX, WD_PARAGRAPH_ALIGNMENT
from typing import Dict

from builder import Attributes, Builder
from hyperlink import add_hyperlink
from table_of_contents import add_table_of_contents
import xmltags


class DocxBuilder(Builder):
    """Docx builder."""

    SCHEMA = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    FORMAT_TAGS = (xmltags.INSTRUCTION, xmltags.BOLD, xmltags.ITALIC, xmltags.STRIKETHROUGH)
    TEXT_TAGS = (
        xmltags.PARAGRAPH, xmltags.LIST_ITEM, xmltags.MEASURE, xmltags.HEADING, xmltags.TABLE_CELL, xmltags.HEADER,
        xmltags.TITLE) + FORMAT_TAGS

    def __init__(self, filename: pathlib.Path, docx_reference_filename: pathlib.Path) -> None:
        super().__init__(filename)
        filename.unlink(missing_ok=True)
        shutil.copy(docx_reference_filename, filename)
        self.doc = Document(filename)
        self.paragraph = None  # The current paragraph
        self.current_list_style = []  # Stack of list styles
        self.previous_list_item = []  # Stack of previous list items
        self.table = None
        self.row = None
        self.column_index = 0
        self.link = None
        self.section_style = None
        self.formatting = set()
        self.style = None

    def start_element(self, tag: str, attributes: Attributes) -> None:
        if tag == xmltags.PARAGRAPH:
            self.paragraph = self.doc.add_paragraph(style=self.style)
        elif tag == xmltags.PAGEBREAK:
            self.doc.add_page_break()
        elif tag == xmltags.BULLET_LIST:
            self.current_list_style.append("Lijst opsom.teken1")
            self.previous_list_item.append(None)
        elif tag == xmltags.NUMBERED_LIST:
            self.current_list_style.append("Lijstnummering1")
            self.previous_list_item.append(None)
        elif tag == xmltags.LIST_ITEM:
            self.paragraph = self.doc.add_paragraph(style=self.current_list_style[-1])
            level = len(self.current_list_style) - 1
            self.paragraph._p.get_or_add_pPr().get_or_add_numPr().get_or_add_ilvl().val = level
            if "Lijstnummering1" == self.current_list_style[-1]:
                if self.previous_list_item[-1] is None:
                    # Add a new concrete numbering for Lijstnummering1. "0" is the id of the abstract numbering of
                    # Lijstnummering1. This id can be found in the word/numbering.xml file (unzip reference.docx so
                    # see word/numbering.xml), look for the <w:abstractNum w:abstractNumId="0" ...> that has a
                    # child element <w:pStyle w:val="Lijstnummering1"/>
                    num = self.doc.part.numbering_part.numbering_definitions._numbering.add_num("0")
                    num.add_lvlOverride(ilvl=level).add_startOverride(1)  # Restart the numbering
                    num = num.numId
                else:
                    num = self.previous_list_item[-1]._p.pPr.numPr.numId.val
                self.paragraph._p.get_or_add_pPr().get_or_add_numPr().get_or_add_numId().val = num
            self.previous_list_item[-1] = self.paragraph
        elif tag == xmltags.SECTION:
            level = attributes[xmltags.SECTION_LEVEL]
            is_appendix = attributes.get(xmltags.SECTION_IS_APPENDIX)
            self.section_style = f"Kop {level} Bijlage" if is_appendix else f"heading {level}"
        elif tag == xmltags.HEADING:
            self.paragraph = self.doc.add_paragraph(style=self.section_style)
        elif tag == xmltags.TABLE:
            self.table = self.doc.add_table(0, int(attributes[xmltags.TABLE_COLUMNS]), style="Tabelraster1")
            # Set table width to 100%
            self.table._tbl.tblPr.xpath("./w:tblW")[0].attrib[f"{self.SCHEMA}type"] = "pct"
            self.table._tbl.tblPr.xpath("./w:tblW")[0].attrib[f"{self.SCHEMA}w"] = "100%"
        elif tag in (xmltags.TABLE_HEADER_ROW, xmltags.TABLE_ROW):
            self.row = self.table.add_row()
            self.column_index = 0
        elif tag == xmltags.TABLE_CELL:
            cell = self.row.cells[self.column_index]
            cell._tc.tcPr.tcW.type = 'auto'
            self.paragraph = cell.paragraphs[0]
            if alignment_attr := attributes.get(xmltags.TABLE_CELL_ALIGNMENT):
                alignment = dict(
                    left=WD_PARAGRAPH_ALIGNMENT.LEFT, right=WD_PARAGRAPH_ALIGNMENT.RIGHT,
                    center=WD_PARAGRAPH_ALIGNMENT.CENTER)[alignment_attr]
                self.paragraph.paragraph_format.alignment = alignment
            self.column_index += 1
        elif tag == xmltags.ANCHOR:
            self.link = attributes[xmltags.ANCHOR_LINK]
        elif tag == xmltags.HEADER:
            self.paragraph = self.doc.sections[0].header.paragraphs[0]
            self.paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        elif tag == xmltags.TITLE:
            self.paragraph = self.doc.add_paragraph(style="Title")
        elif tag == xmltags.TABLE_OF_CONTENTS:
            self.doc.add_paragraph(attributes[xmltags.TABLE_OF_CONTENTS_HEADING], style="TOC Heading")
            add_table_of_contents(self.doc.add_paragraph())
        elif tag == xmltags.MEASURE:
            self.style = "Maatregel"
        elif tag in self.FORMAT_TAGS:
            self.formatting.add(tag)

    def text(self, tag: str, text: str) -> None:
        if tag in self.TEXT_TAGS:
            run = self.paragraph.add_run(text)
            if xmltags.BOLD in self.formatting:
                run.font.bold = True
            if xmltags.INSTRUCTION in self.formatting:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            if xmltags.ITALIC in self.formatting:
                run.font.italic = True
            if xmltags.STRIKETHROUGH in self.formatting:
                run.font.strike = True
        elif tag == xmltags.ANCHOR:
            add_hyperlink(self.paragraph, self.link, text)
        elif tag == xmltags.IMAGE:
            self.doc.add_picture(text)

    def end_element(self, tag: str, attributes: Attributes) -> None:
        if tag in (xmltags.BULLET_LIST, xmltags.NUMBERED_LIST):
            self.current_list_style.pop()
            self.previous_list_item.pop()
        elif tag in self.FORMAT_TAGS:
            self.formatting.remove(tag)
        elif tag == xmltags.MEASURE:
            self.style = None

    def tail(self, tag: str, tail: str, parent: str) -> None:
        self.text(parent, tail)

    def end_document(self) -> None:
        self.doc.save(self.filename)
